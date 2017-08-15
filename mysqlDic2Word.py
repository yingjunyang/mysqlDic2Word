# coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import pymysql

document = Document()
# 加入不同等级的标题
document.add_heading(u'善跑数据库表结构', 0)

# 打开数据库连接
db = pymysql.connect(host='localhost', user='root', password='root', database='information_schema', charset='utf8')
TABLE_SCHEMA = 'test'

# 使用cursor()方法获取操作游标
cursor = db.cursor()

# SQL 查询语句
tb_sql = '''SELECT t.table_name
	,t.table_comment
FROM information_schema.TABLES t
WHERE t.TABLE_SCHEMA = '%s' ''' % (TABLE_SCHEMA)
try:
    # 执行SQL语句
    cursor.execute(tb_sql)
    # 获取所有记录列表
    tbs = cursor.fetchall()
except Exception as err:
    print(err)
    exit(1)
    db.close()

# 增加表格
for row in tbs:
    document.add_heading(row[0], 2)
    paragraph = document.add_paragraph(row[1])
    col_sql = '''SELECT c.column_name
	,c.column_type
	,c.column_key
	,c.is_nullable
	,c.column_comment
    FROM information_schema.COLUMNS c 
    WHERE c.TABLE_SCHEMA = '%s' and c.table_name = '%s' ''' % (TABLE_SCHEMA, row[0])
    try:
        # 执行SQL语句
        cursor.execute(col_sql)
        # 获取所有记录列表
        cols = cursor.fetchall()
    except Exception as err:
        print(err)
        exit(1)
        db.close()
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Key'
    hdr_cells[3].text = 'Is_nullable'
    hdr_cells[4].text = 'Comment'
    for row in cols:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]  # 关闭数据库连接
db.close()
# 增加分页
document.add_page_break()

# 保存文件
document.save(u'表结构.docx')
